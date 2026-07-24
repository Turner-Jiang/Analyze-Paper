"""
文件解析模块 - 支持PDF、DOCX、TXT、图片格式解析
"""

import io
from pathlib import Path

import PyPDF2
from docx import Document
from PIL import Image


def parse_pdf(file_bytes: bytes) -> str:
    """解析PDF文件，提取文本内容"""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """解析DOCX文件，提取文本内容"""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # 也提取表格中的内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_txt(file_bytes: bytes) -> str:
    """解析TXT文件"""
    # 尝试多种编码
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_image(file_bytes: bytes) -> str:
    """
    解析图片文件 - 返回图片描述信息
    注：完整OCR需要额外依赖，这里将图片信息传递给AI进行识别
    """
    img = Image.open(io.BytesIO(file_bytes))
    info = f"[图片文件] 尺寸: {img.size[0]}x{img.size[1]}, 格式: {img.format or '未知'}"
    return info


def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    根据文件扩展名自动选择解析方式
    :param filename: 文件名
    :param file_bytes: 文件内容bytes
    :return: 提取的文本内容
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_bytes)
    elif suffix == ".docx":
        return parse_docx(file_bytes)
    elif suffix == ".txt":
        return parse_txt(file_bytes)
    elif suffix in (".png", ".jpg", ".jpeg"):
        return parse_image(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")


def get_image_base64(file_bytes: bytes) -> str:
    """将图片转为base64，用于AI视觉识别"""
    import base64
    return base64.b64encode(file_bytes).decode("utf-8")
